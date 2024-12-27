import google.generativeai as genai
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import os
from googletrans import Translator
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from collections import deque

# Global API Key will be set after API window input
api_key = None
# Global preference for showing password or not
show_password = False


def validate_api_key(key):
    """Checks to see if the API key looks like a valid Gemini API key."""
    if key and len(key) == 39 and key.startswith("AIzaSy"):
        return True
    return False

def generate_prompt(topic, learning_outcomes, age_group, output_type, time_minutes = None, location_name = None):
     """Generates a prompt based on the input"""
     if output_type == "Ideas":
            prompt = f"""
            Develop a comprehensive and creative set of STEAM (Science, Technology, Engineering, Arts, Mathematics) integration possibilities
            for the topic: "{topic}" tailored exactly for learners aged {age_group}.

            The specified learning outcomes are:
            {', '.join(learning_outcomes)}.

            Provide a detailed description of how to creatively and reliably integrate each of the STEAM disciplines into this topic, emphasizing hands-on
            activities, interdisciplinary projects, and relevant thought-provoking discussions. Include examples and practical explanations that showcase the
            interconnectedness of these areas. Be creative but do not stary too far from the topic. Be realistic. Structure your output using clear headings, sub-headings, and bullet points where necessary, to make it easy to read.
            Use markdown style, for example, use **Heading** for headings and *italic* for italics.
            """
     else:
         prompt = f"""
            Create an optimized and relevant lesson plan for the topic: {topic}
            for learners who are {age_group}. The classroom session is {time_minutes} minutes long.

            The learning outcomes for this lesson are:
            {', '.join(learning_outcomes)}
            
            **Location Context:**
            Location Name: {location_name}

            Given this context, generate a detailed lesson plan which includes STEAM education related integrations.
            The lesson plan should include:

            *   **Title of the Lesson**
            *   **Learning Objectives** (restate main outcomes clearly)
            *   **Materials Needed**
            *   **Detailed Lesson Procedure** (include activity steps and timings)
            *   **STEAM Integration** (explain how Science, Technology, Engineering, Arts, and Mathematics are connected in the activities.)
            *   **Assessment/Review** (check understanding)
            """
     return prompt


class SteamApp:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("STEAM Integration Generator")
        self.window.geometry("1000x700")  # Increased window size
        self.window.configure(bg="#f0f0f0")

        # History tracking
        self.history = deque(maxlen=5)
         # Variables
        self.topic_entry = None
        self.outcomes_entry = None
        self.age_entry = None
        self.language_var = tk.StringVar()
        self.output_text = None
        self.time_entry = None
        self.location_entry = None
        self.output_type_var = tk.StringVar()


        # Styling using ttk
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure("TLabel", font=("Helvetica", 11), background="#f0f0f0")
        self.style.configure("TButton", font=("Helvetica", 11), padding=5)
        self.style.configure("TCombobox", font=("Helvetica", 11))
        self.style.configure("TEntry", font=("Helvetica", 11))
        self.api_key = None # Global api key

        self.create_widgets()

    def generate_steam_ideas(self, topic, learning_outcomes, age_group, output_type, time_minutes = None, location_name = None):
        """Generates elaborate STEAM integration ideas for a topic."""
        global api_key
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-pro")
        prompt = generate_prompt(topic, learning_outcomes, age_group, output_type, time_minutes, location_name)

        try:
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error generating STEAM ideas: {e}"

    def translate_text(self, text, target_language):
        """Translates text using googletrans."""
        translator = Translator()
        try:
            translated = translator.translate(text, dest=target_language)
            return translated.text
        except Exception as e:
            return f"Error in translation: {e}"

    def format_output_text(self, text):
        """Formats the output text with different styles using tags."""
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)

        # Define tags
        self.output_text.tag_config("heading", font=("Helvetica", 14, "bold"))
        self.output_text.tag_config("subheading", font=("Helvetica", 12, "bold"))
        self.output_text.tag_config("italic", font=("Helvetica", 11, "italic"))

        paragraphs = text.split("\n")
        for paragraph in paragraphs:
            words = paragraph.split(" ")
            word_count = 0
            for word in words:
                if word_count < 2 and "**" in word:
                    self.output_text.insert(tk.END, word.replace("**", ""), "heading")
                    word_count = word_count + 1
                elif "**" in word:
                    self.output_text.insert(tk.END, word.replace("**", ""), "subheading")
                    word_count = word_count + 1
                elif "*" in word:
                    self.output_text.insert(tk.END, word.replace("*", ""), "italic")
                    word_count = word_count + 1
                else:
                    self.output_text.insert(tk.END, word + " ")
                    word_count = word_count + 1
            self.output_text.insert(tk.END, "\n")  # Add new line at the end of each paragraph
        self.output_text.config(state=tk.DISABLED)

    def update_history(self, topic, outcomes, age_group, language):
        """Updates the history with current details."""
        history_entry = f"Topic: {topic}, Outcomes: {', '.join(outcomes)}, Age: {age_group}, Lang: {language}"
        self.history.append(history_entry)


    def generate_and_display(self):
        """Handles GUI interactions for STEAM generation and language selection."""
        topic = self.topic_entry.get()
        learning_outcomes = self.outcomes_entry.get().split(",")
        age_group = self.age_entry.get()
        language = self.language_var.get()
        output_type = self.output_type_var.get()
        time_minutes = None
        if output_type == "Lesson Plan":
            time_minutes = self.time_entry.get()
            if not time_minutes:
              messagebox.showerror("Error", "Please provide a value for class time.")
              return
            try:
                time_minutes = int(time_minutes)
            except ValueError:
                messagebox.showerror("Error", "Time must be a valid integer number.")
                return

        location_name = None
        if output_type == "Lesson Plan":
             location_name = self.location_entry.get()

        if not topic or not learning_outcomes or not age_group:
          messagebox.showerror("Error", "Please fill in all fields.")
          return
        steam_ideas = self.generate_steam_ideas(topic, learning_outcomes, age_group, output_type, time_minutes, location_name)


        if steam_ideas.startswith("Error"):
            messagebox.showerror("Error", steam_ideas)
            return

        if language == "Nepali":
            translated_ideas = self.translate_text(steam_ideas, "ne")
            if translated_ideas.startswith("Error"):
                messagebox.showerror("Error", translated_ideas)
                return
            self.format_output_text(translated_ideas)
            self.update_history(topic, learning_outcomes, age_group, "Nepali")

        else:
            self.format_output_text(steam_ideas)
            self.update_history(topic, learning_outcomes, age_group, "English")

    def clear_all(self):
        """Clears all input and output fields."""
        self.topic_entry.delete(0, tk.END)
        self.outcomes_entry.delete(0, tk.END)
        self.age_entry.delete(0, tk.END)
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        self.output_text.config(state=tk.DISABLED)
        if self.time_entry:
          self.time_entry.delete(0,tk.END)
        if self.location_entry:
          self.location_entry.delete(0, tk.END)


    def export_to_docx(self):
        """Exports the generated STEAM ideas to a DOCX file with formatting."""
        steam_text = self.output_text.get(1.0, tk.END).strip()
        if not steam_text:
            messagebox.showwarning("Warning", "No STEAM ideas to export.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word Document", "*.docx")]
        )
        if file_path:
            try:
                document = Document()
                # Apply Helvetica font to the entire document
                style = document.styles["Normal"]
                style.font.name = "Helvetica"
                style.font.size = Pt(11)  # Set default font size

                for paragraph in steam_text.split("\n"):
                    p = document.add_paragraph()
                    words = paragraph.split(" ")
                    word_count = 0
                    for word in words:
                        if word_count < 2 and "**" in word:
                            run = p.add_run(word.replace("**", ""))
                            run.bold = True
                            run.font.size = Pt(14)  # Set bold fonts as headings
                            word_count = word_count + 1
                        elif "**" in word:
                            run = p.add_run(word.replace("**", ""))
                            run.bold = True
                            run.font.size = Pt(13)  # Set bold fonts as sub headings
                            word_count = word_count + 1
                        elif "*" in word:
                            run = p.add_run(word.replace("*", ""))
                            run.italic = True
                            word_count = word_count + 1
                        else:
                            run = p.add_run(word + " ")
                            word_count = word_count + 1

                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                document.save(file_path)
                messagebox.showinfo("Success", "STEAM ideas exported successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export STEAM ideas: {e}")

    def show_about(self):
        """Displays the about message."""
        messagebox.showinfo(
            "About", "STEAM Integration Generator\nCreated by Rajat using Gemini AI.\nhttps://github.com/rajat-ed/steam-integration-generator"
        )

    def show_history(self):
        """Display the history in a message box."""
        if self.history:
            messagebox.showinfo("History of Last 5 Uses", "\n".join(self.history))
        else:
            messagebox.showinfo("History", "No History Found")

    def change_api_key(self):
        """Allows to change the API Key"""
        api_window = tk.Tk()
        api_window.title("Change API Key")
        api_window.geometry("400x150")

        ttk.Label(api_window, text="Enter Your New Gemini API Key:").pack(pady=10)
        api_key_entry = ttk.Entry(api_window, width=50)
        api_key_entry.pack(pady = 10)
        show_password_var = tk.BooleanVar(value = False)
        def toggle_password():
            global show_password
            show_password = not show_password
            if show_password:
                  api_key_entry.config(show="")
            else:
                api_key_entry.config(show="*")

        checkbox = ttk.Checkbutton(api_window, text = "Show password", variable = show_password_var, command = toggle_password)
        checkbox.pack(pady = 5)

        def submit_new_key():
            global api_key
            key = api_key_entry.get()
            if validate_api_key(key):
              api_key = key
              messagebox.showinfo("API Key Changed", "API key changed successfully.")
              api_window.destroy()
            else:
              messagebox.showerror("Error", "Invalid API Key, Please enter a valid key")

        submit_button = ttk.Button(api_window, text="Submit", command=submit_new_key)
        submit_button.pack(pady = 10)
        api_window.mainloop()


    def create_widgets(self):
        """Creates the main application widgets."""
        # Labels and Entries
        ttk.Label(self.window, text="Topic:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
        self.topic_entry = ttk.Entry(self.window, width=80, style="TEntry")  # Increased width
        self.topic_entry.grid(row=0, column=1, padx=10, pady=5)

        ttk.Label(self.window, text="Learning Outcomes (comma separated):").grid(
            row=1, column=0, sticky="w", padx=10, pady=5
        )
        self.outcomes_entry = ttk.Entry(self.window, width=80, style="TEntry")  # Increased width
        self.outcomes_entry.grid(row=1, column=1, padx=10, pady=5)

        ttk.Label(self.window, text="Age Group:").grid(
            row=2, column=0, sticky="w", padx=10, pady=5
        )
        self.age_entry = ttk.Entry(self.window, width=80, style="TEntry")  # Increased width
        self.age_entry.grid(row=2, column=1, padx=10, pady=5)

       # Output type options
        ttk.Label(self.window, text="Output Type:").grid(row=3, column=0, sticky="w", padx=10, pady=5)
        output_options = ["Ideas", "Lesson Plan"]
        output_combobox = ttk.Combobox(self.window, textvariable=self.output_type_var, values = output_options, style = "TCombobox")
        output_combobox.set("Ideas") # Set the default as "Ideas"
        output_combobox.grid(row=3, column=1, padx=10, pady=5)
        # Time entry
        ttk.Label(self.window, text="Class Time (minutes):").grid(
            row=4, column=0, sticky="w", padx=10, pady=5
        )
        self.time_entry = ttk.Entry(self.window, width=80, style="TEntry")  # Increased width
        self.time_entry.grid(row=4, column=1, padx=10, pady=5)

         # Location Input
        ttk.Label(self.window, text="Location Name:").grid(
            row=5, column=0, sticky="w", padx=10, pady=5
        )
        self.location_entry = ttk.Entry(self.window, width=80, style="TEntry")  # Increased width
        self.location_entry.grid(row=5, column=1, padx=10, pady=5)


        # Language Selection
        ttk.Label(self.window, text="Language:").grid(
            row=6, column=0, sticky="w", padx=10, pady=5
        )
        language_combobox = ttk.Combobox(
            self.window,
            textvariable=self.language_var,
            values=["English", "Nepali"],
            style="TCombobox",
        )
        language_combobox.set("English")
        language_combobox.grid(row=6, column=1, padx=10, pady=5)

        # Buttons Frame
        button_frame = ttk.Frame(self.window, padding=10)
        button_frame.grid(row=7, column=0, columnspan=2, pady=10)

        # Generate Button
        generate_button = ttk.Button(
            button_frame, text="Generate STEAM Ideas", command=self.generate_and_display, style="TButton"
        )
        generate_button.grid(row=0, column=0, padx=5)

        # Clear All Button
        clear_button = ttk.Button(
            button_frame, text="Clear All", command=self.clear_all, style="TButton"
        )
        clear_button.grid(row=0, column=1, padx=5)

        # Export Button
        export_button = ttk.Button(
            button_frame, text="Export to Docs", command=self.export_to_docx, style="TButton"
        )
        export_button.grid(row=0, column=2, padx=5)

        # Output Text Box
        self.output_text = scrolledtext.ScrolledText(
            self.window,
            wrap=tk.WORD,
            width=120,  # Increased width
            height=20,  # Increased height
            state=tk.DISABLED,
            font=("Helvetica", 11),
        )
        self.output_text.grid(row=8, column=0, columnspan=2, padx=10, pady=10)

        # Menu Bar
        menu_bar = tk.Menu(self.window)
        self.window.config(menu=menu_bar)

        api_menu = tk.Menu(menu_bar, tearoff = 0)
        api_menu.add_command(label = "Change API Key", command = self.change_api_key)
        menu_bar.add_cascade(label = "API", menu = api_menu)

        history_menu = tk.Menu(menu_bar, tearoff = 0)
        history_menu.add_command(label="History", command=self.show_history)
        menu_bar.add_cascade(label="History", menu = history_menu)

        about_menu = tk.Menu(menu_bar, tearoff=0)
        about_menu.add_command(label="About", command=self.show_about)
        menu_bar.add_cascade(label="About", menu=about_menu)

    def run(self):
        self.window.mainloop()


def api_key_window():
    """Creates a window for API key input."""
    api_window = tk.Tk()
    api_window.title("Enter API Key")
    api_window.geometry("400x150")  # Set api window size

    ttk.Label(api_window, text="Enter Your Gemini API Key:").pack(pady=10)

    api_key_entry = ttk.Entry(api_window, width=50, show="*") # set to * initially
    api_key_entry.pack(pady=10)

    show_password_var = tk.BooleanVar(value = False)
    def toggle_password():
        global show_password
        show_password = not show_password
        if show_password:
              api_key_entry.config(show="")
        else:
            api_key_entry.config(show="*")


    checkbox = ttk.Checkbutton(api_window, text = "Show password", variable = show_password_var, command = toggle_password)
    checkbox.pack(pady = 5)


    def submit_key():
        global api_key
        key = api_key_entry.get()
        if validate_api_key(key):
            api_key = key
            api_window.destroy()  # Close the API window
            app = SteamApp()  # Start the main application
            app.run()
        else:
            messagebox.showerror("Error", "Invalid API Key, Please enter a valid key")

    submit_button = ttk.Button(api_window, text="Submit", command=submit_key)
    submit_button.pack(pady=10)
    api_window.mainloop()

if __name__ == "__main__":
  api_key_window()
